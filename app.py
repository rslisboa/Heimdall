from __future__ import annotations

import json
import os
import re
import unicodedata
from html import unescape
from email.utils import parsedate_to_datetime
from pathlib import Path
from typing import Any
from datetime import datetime, timezone, timedelta
import asyncio
import csv
import io
import uuid
from zoneinfo import ZoneInfo

from urllib.parse import urlparse
import xml.etree.ElementTree as ET

import httpx
from fastapi import FastAPI, Request, WebSocket, WebSocketDisconnect, File, Form, UploadFile
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, Field

from heimdall.agents.voice_assistant import handle_voice_command
from heimdall.db import Database
from heimdall.events import EventBus
from heimdall.models import AgentEvent
from heimdall.settings import Settings, deep_get


# ============================================================
# PAYLOADS
# ============================================================

class VoicePayload(BaseModel):
    text: str


class ChatHistoryItem(BaseModel):
    role: str
    content: str


class ImageAttachment(BaseModel):
    name: str = ""
    type: str = ""
    data_url: str


class LocalLLMRequest(BaseModel):
    pergunta: str
    historico: list[ChatHistoryItem] = Field(default_factory=list)
    imagens: list[ImageAttachment] = Field(default_factory=list)


class LocalActionRequest(LocalLLMRequest):
    executar: bool = True


class TaskCreateRequest(BaseModel):
    title: str
    deadline: str | None = None
    reminder_text: str | None = None
    keywords: list[str] = Field(default_factory=list)
    attachments: list[dict[str, Any]] = Field(default_factory=list)


class TaskUpdateRequest(BaseModel):
    status: str | None = None
    title: str | None = None
    deadline: str | None = None


class LanguageTutorRequest(BaseModel):
    language: str
    message: str
    history: list[ChatHistoryItem] = Field(default_factory=list)


class CoderAnalyzeRequest(BaseModel):
    prompt: str = ""
    file_path: str = ""
    code: str = ""
    language: str = "auto"


# ============================================================
# LLM LOCAL - LM STUDIO
# ============================================================

LMSTUDIO_BASE_URL = os.getenv("LMSTUDIO_BASE_URL", "http://localhost:1234/v1")
LMSTUDIO_MODEL = os.getenv("LMSTUDIO_MODEL", "gemma-4-e2b-it-qat")
LMSTUDIO_TIMEOUT = float(os.getenv("LMSTUDIO_TIMEOUT", "120"))

LANGUAGE_TUTOR_OPTIONS = {
    "en-US": {"name": "ingles", "native": "English"},
    "es-ES": {"name": "espanhol", "native": "espanol"},
    "fr-FR": {"name": "frances", "native": "francais"},
    "ru-RU": {"name": "russo", "native": "русский"},
    "zh-CN": {"name": "chines", "native": "中文"},
    "ja-JP": {"name": "japones", "native": "日本語"},
    "de-DE": {"name": "alemao", "native": "Deutsch"},
    "pt-BR": {"name": "portugues", "native": "portugues brasileiro"},
}

AI_NEWS_KEYWORDS = [
    "inteligencia artificial",
    "inteligencia-artificial",
    "chatgpt",
    "gemini",
    "claude",
    "openai",
    "machine learning",
    "aprendizado de maquina",
    "algoritmo",
]

AI_TOPIC_KEYWORDS = [
    "inteligencia artificial",
    "ia generativa",
    "ias generativas",
    "generative ai",
    "machine learning",
    "aprendizado de maquina",
    "deep learning",
    "large language model",
    "llm",
    "chatgpt",
    "openai",
    "gemini",
    "claude",
    "anthropic",
    "deepseek",
    "agente de ia",
    "agentes de ia",
    "modelos de ia",
    "modelo de ia",
    "redes neurais",
    "neural network",
]

TRUSTED_NEWS_SOURCES = [
    {
        "name": "G1 Tecnologia",
        "feed": "https://g1.globo.com/dynamo/tecnologia/rss2.xml",
        "domains": ["g1.globo.com", "globo.com"],
    },
    {
        "name": "CNN Brasil",
        "feed": "https://www.cnnbrasil.com.br/sitemap-news.xml",
        "domains": ["www.cnnbrasil.com.br", "cnnbrasil.com.br"],
    },
    {
        "name": "Canaltech",
        "feed": "https://canaltech.com.br/rss/",
        "domains": ["canaltech.com.br", "www.canaltech.com.br"],
    },
    {
        "name": "Olhar Digital",
        "feed": "https://olhardigital.com.br/feed/",
        "domains": ["olhardigital.com.br", "www.olhardigital.com.br"],
    },
    {
        "name": "Tecnoblog",
        "feed": "https://tecnoblog.net/feed/",
        "domains": ["tecnoblog.net", "www.tecnoblog.net"],
    },
    {
        "name": "Folha Tecnologia",
        "feed": "https://feeds.folha.uol.com.br/tec/rss091.xml",
        "domains": ["feeds.folha.uol.com.br", "www1.folha.uol.com.br", "folha.uol.com.br"],
    },
    {
        "name": "UOL Tecnologia",
        "feed": "https://rss.uol.com.br/feed/tecnologia.xml",
        "domains": ["rss.uol.com.br", "tecnologia.uol.com.br", "www.uol.com.br", "uol.com.br"],
    },
    {
        "name": "G1 IA",
        "feed": "https://news.google.com/rss/search?q=inteligencia%20artificial%20site%3Ag1.globo.com&hl=pt-BR&gl=BR&ceid=BR:pt-419",
        "domains": ["news.google.com"],
        "source_domains": ["g1.globo.com", "globo.com"],
    },
    {
        "name": "CNN Brasil IA",
        "feed": "https://news.google.com/rss/search?q=inteligencia%20artificial%20site%3Acnnbrasil.com.br&hl=pt-BR&gl=BR&ceid=BR:pt-419",
        "domains": ["news.google.com"],
        "source_domains": ["cnnbrasil.com.br", "www.cnnbrasil.com.br"],
    },
    {
        "name": "Exame IA",
        "feed": "https://news.google.com/rss/search?q=inteligencia%20artificial%20site%3Aexame.com&hl=pt-BR&gl=BR&ceid=BR:pt-419",
        "domains": ["news.google.com"],
        "source_domains": ["exame.com", "www.exame.com"],
    },
    {
        "name": "Folha IA",
        "feed": "https://news.google.com/rss/search?q=inteligencia%20artificial%20site%3Afolha.uol.com.br&hl=pt-BR&gl=BR&ceid=BR:pt-419",
        "domains": ["news.google.com"],
        "source_domains": ["folha.uol.com.br", "www1.folha.uol.com.br"],
    },
]

PODCAST_SOURCES = [
    {
        "name": "IA Sob Controle",
        "feed": "https://www.iasobcontrole.tech/feed/",
        "domains": ["www.iasobcontrole.tech", "iasobcontrole.tech", "open.spotify.com"],
    },
    {
        "name": "Hipsters Ponto Tech",
        "feed": "https://www.hipsters.tech/feed/podcast/",
        "domains": ["www.hipsters.tech", "hipsters.tech", "www.alura.com.br", "alura.com.br"],
    },
    {
        "name": "Data Hackers",
        "feed": "https://anchor.fm/s/6d1ee34/podcast/rss",
        "domains": ["anchor.fm", "open.spotify.com", "www.datahackers.com.br", "datahackers.com.br"],
    },
    {
        "name": "Pizza de Dados",
        "feed": "https://podcast.pizzadedados.com/feed.xml",
        "domains": ["podcast.pizzadedados.com", "pizzadedados.com"],
    },
]

YOUTUBE_VIDEO_SOURCES = [
    {
        "name": "IA Sob Controle",
        "feed": "https://www.youtube.com/feeds/videos.xml?channel_id=UCZgLNCGFiTGnB3XTJIR4g6g",
        "domains": ["www.youtube.com", "youtube.com"],
    },
    {
        "name": "Codigo Fonte TV",
        "feed": "https://www.youtube.com/feeds/videos.xml?channel_id=UCZxr48h7_qEXuM1imy6NcCg",
        "domains": ["www.youtube.com", "youtube.com"],
    },
    {
        "name": "Filipe Deschamps",
        "feed": "https://www.youtube.com/feeds/videos.xml?channel_id=UC70YG2WHVxlOJRng4v-CIFQ",
        "domains": ["www.youtube.com", "youtube.com"],
    },
    {
        "name": "Canaltech",
        "feed": "https://www.youtube.com/feeds/videos.xml?channel_id=UC_bXJnsgwOqEPA_-6N6faKw",
        "domains": ["www.youtube.com", "youtube.com"],
    },
    {
        "name": "Olhar Digital",
        "feed": "https://www.youtube.com/feeds/videos.xml?channel_id=UCGV72aVJuWP0QPNGH4YgIww",
        "domains": ["www.youtube.com", "youtube.com"],
    },
    {
        "name": "Data Hackers",
        "feed": "https://www.youtube.com/feeds/videos.xml?channel_id=UCISrteT3SsMdSkoPWYJ2fGA",
        "domains": ["www.youtube.com", "youtube.com"],
    },
]

PAPER_SOURCES = [
    {
        "name": "arXiv AI + Machine Learning",
        "feed": "https://export.arxiv.org/api/query?search_query=cat:cs.AI+OR+cat:cs.LG+OR+cat:cs.CL+OR+all:%22large%20language%20model%22&sortBy=submittedDate&sortOrder=descending&max_results=16",
        "domains": ["arxiv.org", "export.arxiv.org"],
        "kind": "arxiv",
    },
    {
        "name": "Europe PMC",
        "feed": "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
        "domains": ["europepmc.org", "www.ebi.ac.uk", "doi.org", "www.ncbi.nlm.nih.gov"],
        "kind": "europe_pmc",
    },
    {
        "name": "OpenAlex",
        "feed": "https://api.openalex.org/works",
        "domains": ["openalex.org", "api.openalex.org", "doi.org"],
        "kind": "openalex",
    },
]


def _normalize_text(value: str) -> str:
    normalized = (
        str(value or "")
        .lower()
        .replace("ê", "e")
        .replace("é", "e")
        .replace("è", "e")
        .replace("á", "a")
        .replace("à", "a")
        .replace("ã", "a")
        .replace("â", "a")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("õ", "o")
        .replace("ô", "o")
        .replace("ú", "u")
        .replace("ç", "c")
    )
    return "".join(
        char for char in unicodedata.normalize("NFKD", normalized)
        if not unicodedata.combining(char)
    )


def _safe_domain(url: str, allowed_domains: list[str]) -> bool:
    try:
        parsed = urlparse(url)
    except Exception:
        return False

    if parsed.scheme not in {"http", "https"}:
        return False

    host = parsed.netloc.lower()
    return any(host == domain or host.endswith(f".{domain}") for domain in allowed_domains)


def _xml_text(node: ET.Element, path: str, namespaces: dict[str, str] | None = None) -> str:
    found = node.find(path, namespaces or {})
    return (found.text or "").strip() if found is not None else ""


def _clean_feed_text(value: str, max_length: int = 700) -> str:
    text = str(value or "")
    for _ in range(3):
        text = unescape(text)
    text = re.sub(r"(?is)<script.*?</script>|<style.*?</style>", " ", text)
    text = re.sub(r"(?is)<img\b[^>]*>", " ", text)
    text = re.sub(r"(?is)<a\b[^>]*>(.*?)</a>", r"\1", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = re.sub(r"https?://\S+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:max_length]


def _is_ai_news(title: str, detail: str, link: str) -> bool:
    raw_title_link = f"{title} {link}"
    title_link = _normalize_text(f"{title} {link}")
    detail_head = _normalize_text(str(detail or "")[:360])
    strong_keywords = [
        "inteligencia artificial",
        "chatgpt",
        "gemini",
        "claude",
        "openai",
        "machine learning",
        "aprendizado de maquina",
        "algoritmo",
    ]
    if re.search(r"\b(IA|IAs|AI)\b", raw_title_link):
        return True
    if re.search(r"\b(ia|ias)\b", title_link):
        return True
    if any(keyword in title_link for keyword in strong_keywords):
        return True
    return any(keyword in detail_head for keyword in strong_keywords)


def _is_ai_topic_content(title: str, detail: str, link: str) -> bool:
    raw_title_link = f"{title} {link}"
    title_link = _normalize_text(f"{title} {link}")
    detail_head = _normalize_text(str(detail or "")[:420])
    if re.search(r"\b(IA|IAs|AI)\b", raw_title_link):
        return True
    if re.search(r"\b(ia|ias)\b", title_link):
        return True
    if any(keyword in title_link for keyword in AI_TOPIC_KEYWORDS):
        return True
    detail_keywords = [
        "inteligencia artificial",
        "ia generativa",
        "ias generativas",
        "generative ai",
        "machine learning",
        "aprendizado de maquina",
        "deep learning",
        "large language model",
        "llm",
        "redes neurais",
        "neural network",
    ]
    return any(keyword in detail_head for keyword in detail_keywords)


def _parse_news_timestamp(value: str) -> float:
    if not value:
        return 0.0

    try:
        return parsedate_to_datetime(value).timestamp()
    except Exception:
        pass

    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).timestamp()
    except Exception:
        return 0.0


def _is_current_item(item: dict[str, Any], days: int = 15) -> bool:
    timestamp = float(item.get("sort_timestamp") or 0)
    if timestamp <= 0:
        return False
    cutoff = datetime.now(timezone.utc).timestamp() - (days * 24 * 60 * 60)
    return timestamp >= cutoff


def _is_current_month_item(item: dict[str, Any]) -> bool:
    timestamp = float(item.get("sort_timestamp") or 0)
    if timestamp <= 0:
        return False
    now = datetime.now(timezone.utc)
    month_start = datetime(now.year, now.month, 1, tzinfo=timezone.utc).timestamp()
    return timestamp >= month_start


def _filter_recent_or_month(items: list[dict[str, Any]], limit: int) -> list[dict[str, Any]]:
    ordered = sorted(
        items,
        key=lambda item: float(item.get("sort_timestamp") or 0),
        reverse=True,
    )
    recent = [item for item in ordered if _is_current_item(item, 15)]
    if len(recent) >= limit:
        return recent
    current_month = [item for item in ordered if _is_current_month_item(item)]
    return current_month or recent


def _fallback_news_rating(item: dict[str, Any]) -> dict[str, Any]:
    text = _normalize_text(f"{item.get('title', '')} {item.get('detail', '')}")
    stars = 2
    sentiment = "neutro"
    reasons: list[str] = []

    if item.get("type") == "paper":
        if any(term in text for term in ["large language model", "language model", "llm", "agentic", "agente", "multimodal"]):
            stars += 2
            reasons.append("paper tecnico diretamente ligado a LLMs, agentes ou modelos multimodais")
        elif any(term in text for term in ["artificial intelligence", "inteligencia artificial", "machine learning", "deep learning"]):
            stars += 1
            reasons.append("paper academico recente sobre IA ou machine learning")
        if any(term in text for term in ["review", "bibliometric", "narrative review", "survey"]):
            stars = min(stars, 3)
            reasons.append("conteudo de revisao, util mas menos acionavel")
        return {
            "stars": max(1, min(5, stars)),
            "sentiment": "neutro",
            "rating_reason": "; ".join(reasons[:2]) or "documento tecnico recente indexado em fonte academica confiavel",
        }

    if any(term in text for term in ["openai", "gemini", "chatgpt", "claude", "grok"]):
        stars += 1
        reasons.append("envolve modelo ou empresa central de IA")

    if any(term in text for term in ["regulacao", "seguranca", "falha", "risco", "justica", "governo"]):
        stars += 1
        reasons.append("tem impacto regulatorio, juridico ou de seguranca")

    if any(term in text for term in ["brasil", "gratuita", "empresa", "mercado", "produtividade", "negocio"]):
        stars += 1
        reasons.append("tem aplicacao pratica para negocios ou Brasil")

    if any(term in text for term in ["falha", "risco", "acusada", "ilegal", "vicio", "seguranca"]):
        sentiment = "alerta"
    elif any(term in text for term in ["gratuita", "melhor", "inovacao", "lancamento"]):
        sentiment = "positivo"

    return {
        "stars": max(1, min(5, stars)),
        "sentiment": sentiment,
        "rating_reason": "; ".join(reasons[:2]) or "relevancia moderada por contexto de tecnologia e IA",
    }


def _extract_json_from_llm(content: str) -> Any:
    cleaned = str(content or "").strip()

    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()

    try:
        return json.loads(cleaned)
    except Exception:
        pass

    decoder = json.JSONDecoder()

    for start_char in ("{", "["):
        start = cleaned.find(start_char)

        while start >= 0:
            try:
                parsed, _ = decoder.raw_decode(cleaned[start:])
                return parsed
            except Exception:
                start = cleaned.find(start_char, start + 1)

    raise ValueError("Resposta da LLM nao contem JSON valido.")


async def classificar_noticias_com_llm(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not items:
        return []

    news_payload = [
        {
            "id": idx,
            "fonte": item.get("source", ""),
            "titulo": item.get("title", ""),
            "resumo": item.get("detail", "")[:450],
        }
        for idx, item in enumerate(items[:24])
    ]

    prompt = (
        "Classifique noticias sobre inteligencia artificial para o painel Heimdall.\n"
        "Para cada item, retorne JSON estrito em uma lista chamada items.\n"
        "Campos: id, stars inteiro de 1 a 5, sentiment em positivo|neutro|alerta|critico, reason curta.\n"
        "Criterios de estrelas: impacto no Brasil/negocios/seguranca/regulacao/produtividade, novidade e relevancia pratica.\n\n"
        f"Noticias:\n{json.dumps(news_payload, ensure_ascii=False)}"
    )

    payload = {
        "model": LMSTUDIO_MODEL,
        "messages": [
            {
                "role": "system",
                "content": "Voce e um classificador editorial. Responda apenas JSON valido.",
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.1,
        "max_tokens": 1400,
        "stream": False,
    }

    ratings: dict[int, dict[str, Any]] = {}

    try:
        async with httpx.AsyncClient(timeout=min(LMSTUDIO_TIMEOUT, 45)) as client:
            response = await client.post(f"{LMSTUDIO_BASE_URL}/chat/completions", json=payload)
        response.raise_for_status()
        data = response.json()
        content = str(data.get("choices", [{}])[0].get("message", {}).get("content") or "{}").strip()

        if content.startswith("```"):
            content = content.strip("`").replace("json", "", 1).strip()

        parsed = _extract_json_from_llm(content)
        parsed_items = parsed.get("items", parsed if isinstance(parsed, list) else [])

        for rating in parsed_items:
            idx = int(rating.get("id"))
            stars = int(rating.get("stars", 3))
            sentiment = str(rating.get("sentiment", "neutro")).lower()
            if sentiment not in {"positivo", "neutro", "alerta", "critico"}:
                sentiment = "neutro"
            ratings[idx] = {
                "stars": max(1, min(5, stars)),
                "sentiment": sentiment,
                "rating_reason": str(rating.get("reason", ""))[:220],
            }

    except Exception:
        ratings = {}

    classified: list[dict[str, Any]] = []
    for idx, item in enumerate(items):
        fallback_rating = _fallback_news_rating(item)
        rating = ratings.get(idx)
        if rating:
            llm_stars = int(rating.get("stars") or 3)
            fallback_stars = int(fallback_rating.get("stars") or 3)
            rating["stars"] = max(1, min(5, int((llm_stars + fallback_stars) / 2)))
            rating["sentiment"] = rating.get("sentiment") or fallback_rating.get("sentiment")
            rating["rating_reason"] = rating.get("rating_reason") or fallback_rating.get("rating_reason")
        else:
            rating = fallback_rating
        classified.append({**item, **rating})

    return classified


def diversificar_noticias(items: list[dict[str, Any]], limit: int) -> list[dict[str, Any]]:
    if not items:
        return []

    bounded_limit = max(1, min(limit, 30))
    by_source: dict[str, list[dict[str, Any]]] = {}

    for item in items:
        by_source.setdefault(str(item.get("source") or "Fonte"), []).append(item)

    for source_items in by_source.values():
        source_items.sort(
            key=lambda x: (
                int(x.get("stars") or 0),
                float(x.get("sort_timestamp") or 0),
            ),
            reverse=True,
        )

    source_names = sorted(
        by_source,
        key=lambda name: (
            int(by_source[name][0].get("stars") or 0),
            float(by_source[name][0].get("sort_timestamp") or 0),
        ),
        reverse=True,
    )

    selected: list[dict[str, Any]] = []
    used_urls: set[str] = set()

    while len(selected) < bounded_limit and any(by_source.values()):
        advanced = False

        for source_name in source_names:
            bucket = by_source.get(source_name) or []

            while bucket:
                candidate = bucket.pop(0)
                url = str(candidate.get("url") or candidate.get("title") or "")

                if url in used_urls:
                    continue

                used_urls.add(url)
                selected.append(candidate)
                advanced = True
                break

            if len(selected) >= bounded_limit:
                break

        if not advanced:
            break

    selected.sort(
        key=lambda x: (
            int(x.get("stars") or 0),
            float(x.get("sort_timestamp") or 0),
        ),
        reverse=True,
    )

    selected = selected[:bounded_limit]
    star_values = [int(item.get("stars") or 0) for item in selected]
    if len(selected) >= 3 and len(set(star_values)) == 1:
        base = max(1, min(5, star_values[0]))
        for idx, item in enumerate(selected):
            if idx < 2:
                item["stars"] = base
            elif idx < 4:
                item["stars"] = max(1, base - 1)
            else:
                item["stars"] = max(1, base - 2)

    return selected


def _dedupe_feed_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    deduped: list[dict[str, Any]] = []
    seen: set[str] = set()

    for item in items:
        key = str(item.get("url") or item.get("pdf_url") or item.get("title") or "")
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(item)

    return deduped


async def _classify_and_rank_items(items: list[dict[str, Any]], limit: int) -> list[dict[str, Any]]:
    deduped = _dedupe_feed_items(_filter_recent_or_month(items, max(1, min(limit, 5))))
    classified = await classificar_noticias_com_llm(deduped)
    classified.sort(
        key=lambda x: (
            int(x.get("stars") or 0),
            float(x.get("sort_timestamp") or 0),
        ),
        reverse=True,
    )
    return diversificar_noticias(classified, max(1, min(limit, 5)))


async def buscar_noticias_ia(limit: int = 12) -> dict[str, Any]:
    """
    Busca noticias sobre IA apenas em fontes explicitamente permitidas.
    Usa RSS/sitemap publico, sem chave, e descarta links fora da lista branca.
    """
    section_limit = max(1, min(limit, 5))
    items: list[dict[str, Any]] = []
    source_status: list[dict[str, Any]] = []

    async with httpx.AsyncClient(timeout=18, follow_redirects=True) as client:
        for source in TRUSTED_NEWS_SOURCES:
            try:
                response = await client.get(source["feed"])
                response.raise_for_status()
                root = ET.fromstring(response.content)
            except Exception as exc:
                source_status.append({"name": source["name"], "ok": False, "error": str(exc)})
                continue

            source_status.append({"name": source["name"], "ok": True, "error": ""})

            if root.tag.endswith("rss"):
                for item in root.findall("./channel/item"):
                    title = _clean_feed_text(_xml_text(item, "title"), 240)
                    detail = _clean_feed_text(_xml_text(item, "description"), 620)
                    link = _xml_text(item, "link")
                    published_at = _xml_text(item, "pubDate")
                    original_source = item.find("source")
                    original_source_name = (original_source.text or "").strip() if original_source is not None else ""
                    original_source_url = original_source.attrib.get("url", "") if original_source is not None else ""

                    if not link or not _safe_domain(link, source["domains"]):
                        continue

                    if source.get("source_domains"):
                        if not original_source_url or not _safe_domain(original_source_url, source["source_domains"]):
                            continue

                    if not _is_ai_news(title, detail, link):
                        continue

                    items.append(
                        {
                            "source": original_source_name or source["name"],
                            "title": title,
                            "detail": detail,
                            "url": link,
                            "published_at": published_at,
                            "sort_timestamp": _parse_news_timestamp(published_at),
                            "trusted": True,
                            "source_url": original_source_url,
                        }
                    )

            else:
                namespaces = {
                    "news": "http://www.google.com/schemas/sitemap-news/0.9",
                    "sm": "http://www.sitemaps.org/schemas/sitemap/0.9",
                }

                for url_node in root.findall(".//sm:url", namespaces):
                    link = _xml_text(url_node, "sm:loc", namespaces)
                    title = _clean_feed_text(_xml_text(url_node, "news:news/news:title", namespaces), 240)
                    published_at = _xml_text(url_node, "news:news/news:publication_date", namespaces)

                    if not link or not _safe_domain(link, source["domains"]):
                        continue

                    if "/tecnologia/" not in link and not _is_ai_news(title, "", link):
                        continue

                    if not _is_ai_news(title, "", link):
                        continue

                    items.append(
                        {
                            "source": source["name"],
                            "title": title,
                            "detail": "",
                            "url": link,
                            "published_at": published_at,
                            "sort_timestamp": _parse_news_timestamp(published_at),
                            "trusted": True,
                        }
                    )

    diversified = await _classify_and_rank_items(items, section_limit)
    podcasts = await buscar_podcasts_ia(limit=section_limit)
    videos = await buscar_videos_ia(limit=section_limit)
    papers = await buscar_papers_ia(limit=section_limit)

    sections = {
        "news": {
            "title": "Noticias de IA",
            "subtitle": "Portais confiaveis e variados, com classificacao da IA local.",
            "items": diversified,
        },
        "podcasts": {
            "title": "Podcasts sobre IA",
            "subtitle": "Curadoria brasileira com IA Sob Controle, Hipsters, Data Hackers e Pizza de Dados.",
            "items": podcasts.get("items", []),
        },
        "videos": {
            "title": "Videos no YouTube",
            "subtitle": "Canais brasileiros de tecnologia filtrados por temas de IA.",
            "items": videos.get("items", []),
        },
        "papers": {
            "title": "Documentos e Papers",
            "subtitle": "Papers tecnicos recentes com link para leitura e PDF quando disponivel.",
            "items": papers.get("items", []),
        },
    }

    return {
        "ok": True,
        "checked_at": datetime.now(timezone.utc).isoformat(),
        "sources": [
            {"name": source["name"], "feed": source["feed"], "domains": source["domains"]}
            for source in TRUSTED_NEWS_SOURCES
        ],
        "source_status": source_status + podcasts.get("source_status", []) + videos.get("source_status", []) + papers.get("source_status", []),
        "sections": sections,
        "items": diversified,
    }


async def buscar_podcasts_ia(limit: int = 8) -> dict[str, Any]:
    items: list[dict[str, Any]] = []
    source_status: list[dict[str, Any]] = []

    async with httpx.AsyncClient(timeout=18, follow_redirects=True, headers={"User-Agent": "Heimdall/1.0"}) as client:
        for source in PODCAST_SOURCES:
            try:
                response = await client.get(source["feed"])
                response.raise_for_status()
                root = ET.fromstring(response.content)
            except Exception as exc:
                source_status.append({"name": source["name"], "type": "podcast", "ok": False, "error": str(exc)})
                continue

            source_status.append({"name": source["name"], "type": "podcast", "ok": True, "error": ""})

            for item in root.findall("./channel/item")[:18]:
                title = _clean_feed_text(_xml_text(item, "title"), 240)
                detail = _clean_feed_text(
                    _xml_text(item, "description") or _xml_text(item, "{http://purl.org/rss/1.0/modules/content/}encoded"),
                    620,
                )
                link = _xml_text(item, "link") or source["feed"]
                published_at = _xml_text(item, "pubDate")

                if not _is_ai_topic_content(title, detail, link):
                    continue

                items.append(
                    {
                        "type": "podcast",
                        "source": source["name"],
                        "title": title,
                        "detail": detail,
                        "url": link,
                        "published_at": published_at,
                        "sort_timestamp": _parse_news_timestamp(published_at),
                        "trusted": True,
                    }
                )

    ranked = await _classify_and_rank_items(items, limit)
    return {"ok": True, "source_status": source_status, "items": ranked}


async def buscar_videos_ia(limit: int = 8) -> dict[str, Any]:
    items: list[dict[str, Any]] = []
    source_status: list[dict[str, Any]] = []
    namespaces = {
        "atom": "http://www.w3.org/2005/Atom",
        "media": "http://search.yahoo.com/mrss/",
    }

    async with httpx.AsyncClient(timeout=18, follow_redirects=True, headers={"User-Agent": "Heimdall/1.0"}) as client:
        for source in YOUTUBE_VIDEO_SOURCES:
            try:
                response = await client.get(source["feed"])
                response.raise_for_status()
                root = ET.fromstring(response.content)
            except Exception as exc:
                source_status.append({"name": source["name"], "type": "video", "ok": False, "error": str(exc)})
                continue

            source_status.append({"name": source["name"], "type": "video", "ok": True, "error": ""})

            for entry in root.findall("atom:entry", namespaces)[:12]:
                title = _clean_feed_text(_xml_text(entry, "atom:title", namespaces), 240)
                detail = _clean_feed_text(_xml_text(entry, "media:group/media:description", namespaces), 620)
                published_at = _xml_text(entry, "atom:published", namespaces)
                link_node = entry.find("atom:link[@rel='alternate']", namespaces)
                link = link_node.attrib.get("href", "") if link_node is not None else ""

                if not link or not _safe_domain(link, source["domains"]):
                    continue

                if not _is_ai_topic_content(title, detail, link):
                    continue

                items.append(
                    {
                        "type": "video",
                        "source": source["name"],
                        "title": title,
                        "detail": detail,
                        "url": link,
                        "published_at": published_at,
                        "sort_timestamp": _parse_news_timestamp(published_at),
                        "trusted": True,
                    }
                )

    ranked = await _classify_and_rank_items(items, limit)
    return {"ok": True, "source_status": source_status, "items": ranked}


async def buscar_papers_ia(limit: int = 8) -> dict[str, Any]:
    items: list[dict[str, Any]] = []
    source_status: list[dict[str, Any]] = []
    namespaces = {
        "atom": "http://www.w3.org/2005/Atom",
        "arxiv": "http://arxiv.org/schemas/atom",
    }
    now = datetime.now(timezone.utc)
    month_start = datetime(now.year, now.month, 1, tzinfo=timezone.utc)
    date_start = (now - timedelta(days=15)).date().isoformat()
    date_end = now.date().isoformat()

    async with httpx.AsyncClient(timeout=18, follow_redirects=True, headers={"User-Agent": "Heimdall/1.0"}) as client:
        for source in PAPER_SOURCES:
            try:
                kind = source.get("kind", "arxiv")
                if kind == "europe_pmc":
                    response = await client.get(
                        source["feed"],
                        params={
                            "query": f'("artificial intelligence" OR "machine learning" OR "large language model") FIRST_PDATE:[{month_start.date().isoformat()} TO {date_end}]',
                            "format": "json",
                            "pageSize": "16",
                            "sort_date": "y",
                            "resultType": "core",
                        },
                    )
                elif kind == "openalex":
                    response = await client.get(
                        source["feed"],
                        params={
                            "search": "artificial intelligence machine learning large language model",
                            "filter": f"from_publication_date:{month_start.date().isoformat()},to_publication_date:{date_end},type:article",
                            "sort": "publication_date:desc",
                            "per-page": "16",
                        },
                    )
                else:
                    response = await client.get(source["feed"])
                response.raise_for_status()
            except Exception as exc:
                source_status.append({"name": source["name"], "type": "paper", "ok": False, "error": str(exc)})
                continue

            source_status.append({"name": source["name"], "type": "paper", "ok": True, "error": ""})

            kind = source.get("kind", "arxiv")

            if kind == "europe_pmc":
                data = response.json()
                for paper in data.get("resultList", {}).get("result", [])[:16]:
                    title = _clean_feed_text(paper.get("title", ""), 260)
                    summary = _clean_feed_text(paper.get("abstractText", ""), 760)
                    published_at = paper.get("firstPublicationDate") or paper.get("firstIndexDate") or ""
                    link = ""
                    pdf_url = ""
                    full_urls = paper.get("fullTextUrlList", {}).get("fullTextUrl", [])
                    for url_item in full_urls:
                        candidate = url_item.get("url", "")
                        style = str(url_item.get("documentStyle", "")).lower()
                        availability = str(url_item.get("availabilityCode", "")).upper()
                        if style == "pdf" and availability in {"OA", "F"}:
                            pdf_url = candidate
                        if not link and candidate:
                            link = candidate
                    if not link and paper.get("doi"):
                        link = f"https://doi.org/{paper['doi']}"
                    if not link and paper.get("pmcid"):
                        link = f"https://europepmc.org/articles/{paper['pmcid']}"
                    if not link or not _safe_domain(link, source["domains"]):
                        continue
                    items.append(
                        {
                            "type": "paper",
                            "source": source["name"],
                            "title": title,
                            "detail": summary,
                            "url": link,
                            "pdf_url": pdf_url,
                            "authors": _clean_feed_text(paper.get("authorString", ""), 180),
                            "published_at": published_at,
                            "sort_timestamp": _parse_news_timestamp(published_at),
                            "trusted": True,
                            "technical_document": True,
                        }
                    )
                continue

            if kind == "openalex":
                data = response.json()
                for paper in data.get("results", [])[:16]:
                    location = paper.get("primary_location") or {}
                    open_access = paper.get("open_access") or {}
                    title = _clean_feed_text(paper.get("display_name") or paper.get("title") or "", 260)
                    published_at = paper.get("publication_date") or ""
                    link = location.get("landing_page_url") or open_access.get("oa_url") or paper.get("doi") or paper.get("id") or ""
                    pdf_url = location.get("pdf_url") or ""
                    authors = ", ".join(
                        _clean_feed_text((author.get("author") or {}).get("display_name", ""), 80)
                        for author in paper.get("authorships", [])[:4]
                    )
                    source_name = ((location.get("source") or {}).get("display_name") or "").strip()
                    detail = source_name or "Registro academico indexado pelo OpenAlex."
                    if not link or not _safe_domain(link, source["domains"]):
                        continue
                    items.append(
                        {
                            "type": "paper",
                            "source": source["name"],
                            "title": title,
                            "detail": detail,
                            "url": link,
                            "pdf_url": pdf_url,
                            "authors": authors,
                            "published_at": published_at,
                            "sort_timestamp": _parse_news_timestamp(published_at),
                            "trusted": True,
                            "technical_document": True,
                        }
                    )
                continue

            root = ET.fromstring(response.content)
            for entry in root.findall("atom:entry", namespaces)[:16]:
                    title = _clean_feed_text(_xml_text(entry, "atom:title", namespaces), 260)
                    summary = _clean_feed_text(_xml_text(entry, "atom:summary", namespaces), 760)
                    published_at = _xml_text(entry, "atom:published", namespaces)
                    link = _xml_text(entry, "atom:id", namespaces)
                    pdf_url = ""
                    authors = [
                        _clean_feed_text(_xml_text(author, "atom:name", namespaces), 80)
                        for author in entry.findall("atom:author", namespaces)[:4]
                    ]

                    for link_node in entry.findall("atom:link", namespaces):
                        if link_node.attrib.get("title") == "pdf":
                            pdf_url = link_node.attrib.get("href", "")
                        elif link_node.attrib.get("rel") == "alternate":
                            link = link_node.attrib.get("href", link)

                    if not link or not _safe_domain(link, source["domains"]):
                        continue

                    items.append(
                        {
                            "type": "paper",
                            "source": source["name"],
                            "title": title,
                            "detail": summary,
                            "url": link,
                            "pdf_url": pdf_url,
                            "authors": ", ".join([author for author in authors if author]),
                            "published_at": published_at,
                            "sort_timestamp": _parse_news_timestamp(published_at),
                            "trusted": True,
                            "technical_document": True,
                        }
                    )

    ranked = await _classify_and_rank_items(items, limit)
    return {"ok": True, "source_status": source_status, "items": ranked}


async def testar_conexao_lmstudio() -> dict:
    """
    Testa se o LM Studio Local Server está ativo.
    """
    url = f"{LMSTUDIO_BASE_URL}/models"

    async with httpx.AsyncClient(timeout=10) as client:
        response = await client.get(url)

    response.raise_for_status()
    return response.json()


async def perguntar_llm_local(
    pergunta: str,
    contexto_operacional: str = "",
    historico: list[ChatHistoryItem] | None = None,
    imagens: list[ImageAttachment] | None = None,
) -> str:
    """
    Envia uma pergunta para o modelo local carregado no LM Studio.
    Aceita contexto operacional do Heimdall, histórico de conversa e imagens.
    """
    if not pergunta or not pergunta.strip():
        return "Pergunta vazia."

    url = f"{LMSTUDIO_BASE_URL}/chat/completions"
    pergunta_limpa = pergunta.strip()
    historico = historico or []
    imagens = imagens or []

    contexto_heimdall = (
        "Contexto opcional sobre o sistema Heimdall:\n"
        "Heimdall é um sistema local de monitoramento pessoal criado em Python, "
        "com dashboard FastAPI, banco SQLite, eventos em tempo real, agentes para e-mail, "
        "calendário, voz e integrações futuras. Ele funciona como um painel central para "
        "acompanhar informações importantes e acionar automações locais.\n\n"
    )

    contexto_para_prompt = (
        contexto_operacional
        if contexto_operacional
        else "Nenhum contexto operacional foi carregado para esta pergunta."
    )

    async def chamar_modelo(prompt_usuario: str, max_tokens: int = 1024) -> dict:
        messages: list[dict[str, Any]] = [
            {
                "role": "system",
                "content": (
                    "Você é a IA local do sistema Heimdall. "
                    "Você funciona como um assistente profissional de conversa, análise e apoio operacional. "
                    "Responda sempre em português brasileiro, com clareza, objetividade e contexto."
                ),
            }
        ]

        for item in historico[-10:]:
            role = item.role if item.role in {"user", "assistant"} else "user"
            content = str(item.content or "").strip()

            # Evita repetir a pergunta atual quando o front já a colocou no histórico.
            if not content or (role == "user" and content == pergunta_limpa):
                continue

            messages.append(
                {
                    "role": role,
                    "content": content,
                }
            )

        if imagens:
            user_content: list[dict[str, Any]] = [
                {
                    "type": "text",
                    "text": (
                        prompt_usuario
                        + "\n\nO usuário anexou uma ou mais imagens. "
                        + "Analise visualmente as imagens se o modelo carregado tiver suporte a visão. "
                        + "Se o modelo atual não conseguir interpretar imagem, informe isso claramente."
                    ),
                }
            ]

            for img in imagens[:4]:
                user_content.append(
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": img.data_url,
                        },
                    }
                )

            messages.append(
                {
                    "role": "user",
                    "content": user_content,
                }
            )

        else:
            messages.append(
                {
                    "role": "user",
                    "content": prompt_usuario,
                }
            )

        payload = {
            "model": LMSTUDIO_MODEL,
            "messages": messages,
            "temperature": 0.4,
            "top_p": 0.9,
            "max_tokens": max_tokens,
            "stream": False,
        }

        async with httpx.AsyncClient(timeout=LMSTUDIO_TIMEOUT) as client:
            response = await client.post(url, json=payload)

        if response.status_code >= 400:
            if imagens:
                return {
                    "erro": (
                        "Recebi a imagem, mas o modelo carregado no LM Studio pode não ter suporte visual "
                        f"ou recusou o formato enviado. HTTP {response.status_code} - {response.text}"
                    )
                }

            return {
                "erro": f"Erro ao chamar LM Studio: HTTP {response.status_code} - {response.text}"
            }

        return response.json()

    def extrair_resposta(data: dict) -> str:
        if "erro" in data:
            return data["erro"]

        choices = data.get("choices", [])

        if not choices:
            return ""

        choice = choices[0]
        message = choice.get("message", {})

        content = (
            message.get("content")
            or message.get("reasoning_content")
            or choice.get("text")
            or ""
        )

        return str(content).strip()

    def resposta_ruim(texto: str) -> bool:
        if not texto:
            return True

        texto_limpo = texto.strip()

        return texto_limpo in {
            ".",
            "..",
            "...",
            "-",
            "_",
            ".**",
            "**",
        }

    prompt_1 = (
        "Você é uma IA local de uso geral rodando dentro da interface do Heimdall.\n"
        "O Heimdall é o sistema onde você está integrado.\n\n"
        f"{contexto_heimdall}"
        "Contexto operacional disponível neste momento:\n"
        f"{contexto_para_prompt}\n\n"
        "Regras de resposta:\n"
        "1. Se a pergunta for sobre e-mails, reuniões, eventos, Slack, tarefas ou dados do Heimdall, use o contexto operacional acima.\n"
        "2. Se existir contexto operacional suficiente, nunca diga que não tem acesso aos dados.\n"
        "3. Se o contexto operacional trouxer uma lista, considere o item 1 como o mais recente.\n"
        "4. Se o contexto operacional não tiver a informação solicitada, explique exatamente o que está faltando.\n"
        "5. Se a pergunta for geral, responda normalmente usando seu conhecimento interno.\n"
        "6. Se houver imagens anexadas, tente analisá-las visualmente.\n"
        "7. Se o modelo atual não suportar visão, diga claramente que a imagem foi recebida, mas o modelo carregado não consegue interpretá-la.\n"
        "8. Responda sempre em português brasileiro, de forma clara, objetiva e útil.\n\n"
        f"Pergunta do usuário:\n{pergunta_limpa}"
    )

    data_1 = await chamar_modelo(prompt_1, max_tokens=1024)
    resposta_1 = extrair_resposta(data_1)

    if not resposta_ruim(resposta_1):
        return resposta_1

    prompt_2 = (
        "A resposta anterior veio vazia, inválida ou restritiva demais.\n\n"
        "Você é uma IA local de uso geral dentro do Heimdall.\n"
        "Responda a pergunta do usuário de forma direta.\n"
        "Use o contexto operacional se a pergunta for sobre dados do Heimdall.\n"
        "Se houver contexto suficiente, não diga que não tem acesso aos dados.\n"
        "Para perguntas gerais, responda normalmente com seu conhecimento interno.\n"
        "Se houver imagens anexadas, tente analisá-las se o modelo atual tiver suporte visual.\n\n"
        "Contexto operacional:\n"
        f"{contexto_para_prompt}\n\n"
        f"Pergunta:\n{pergunta_limpa}"
    )

    data_2 = await chamar_modelo(prompt_2, max_tokens=768)
    resposta_2 = extrair_resposta(data_2)

    if not resposta_ruim(resposta_2):
        return resposta_2

    prompt_3 = (
        "Responda em português brasileiro, de forma simples e direta.\n\n"
        "Contexto disponível:\n"
        f"{contexto_para_prompt}\n\n"
        f"Pergunta:\n{pergunta_limpa}"
    )

    data_3 = await chamar_modelo(prompt_3, max_tokens=512)
    resposta_3 = extrair_resposta(data_3)

    if not resposta_ruim(resposta_3):
        return resposta_3

    return (
        "A LLM local foi chamada corretamente, mas o modelo retornou resposta vazia "
        "ou apenas pontuação em três tentativas. Isso indica instabilidade do modelo "
        "local nessa configuração."
    )


# ============================================================
# APP FACTORY
# ============================================================

def create_app(*, db: Database, bus: EventBus, settings: Settings, config: dict[str, Any]) -> FastAPI:
    app = FastAPI(title="Heimdall Dashboard")

    templates_path = Path(__file__).parent / "templates"
    templates = Jinja2Templates(directory=str(templates_path))

    def valor_evento(evento: Any, *nomes: str, default: str = "") -> str:
        """
        Lê valores de eventos vindos como dict, sqlite.Row ou objeto.
        """
        for nome in nomes:
            try:
                if isinstance(evento, dict) and nome in evento:
                    valor = evento.get(nome)
                    if valor is not None:
                        return str(valor)

                valor = getattr(evento, nome, None)
                if valor is not None:
                    return str(valor)

                try:
                    valor = evento[nome]
                    if valor is not None:
                        return str(valor)
                except Exception:
                    pass

            except Exception:
                pass

        return default

    def metadata_evento(evento: Any) -> dict[str, Any]:
        """
        Retorna metadata do evento como dict, quando disponível.
        """
        raw = None

        try:
            if isinstance(evento, dict):
                raw = evento.get("metadata")
            else:
                raw = getattr(evento, "metadata", None)
        except Exception:
            raw = None

        if raw is None:
            try:
                raw = evento["metadata"]
            except Exception:
                raw = None

        if isinstance(raw, dict):
            return raw

        if isinstance(raw, str) and raw.strip():
            try:
                data = json.loads(raw)
                return data if isinstance(data, dict) else {}
            except Exception:
                return {}

        return {}

    async def montar_contexto_operacional(pergunta: str) -> str:
        """
        Monta contexto do Heimdall para perguntas sobre dados internos.
        """
        pergunta_norm = (pergunta or "").lower()

        precisa_email = any(
            palavra in pergunta_norm
            for palavra in [
                "email",
                "e-mail",
                "e-mails",
                "último email",
                "ultimo email",
                "último e-mail",
                "ultimo e-mail",
                "mensagem recebida",
                "caixa de entrada",
                "inbox",
                "recebi",
                "recebido",
            ]
        )

        precisa_reuniao = any(
            palavra in pergunta_norm
            for palavra in [
                "reunião",
                "reuniao",
                "reuniões",
                "reunioes",
                "agenda",
                "calendário",
                "calendario",
                "evento",
                "eventos",
            ]
        )

        precisa_slack = "slack" in pergunta_norm
        precisa_noticia = any(
            palavra in pergunta_norm
            for palavra in ["noticia", "noticias", "noticia", "noticias", "ia", "inteligencia artificial"]
        )
        precisa_tarefa = any(
            palavra in pergunta_norm
            for palavra in ["tarefa", "tarefas", "task", "tasks", "pendência", "pendencia"]
        )

        if precisa_noticia:
            try:
                news = await buscar_noticias_ia(limit=8)
                linhas_news = [
                    "Noticias confiaveis recentes sobre IA, filtradas por dominio permitido:",
                ]

                for idx, item in enumerate(news.get("items", [])[:8], start=1):
                    linhas_news.append(
                        f"{idx}. Fonte: {item.get('source', '')}\n"
                        f"   Titulo: {item.get('title', '')}\n"
                        f"   Data: {item.get('published_at', '')}\n"
                        f"   Link: {item.get('url', '')}"
                    )

                if len(linhas_news) > 1:
                    return "\n".join(linhas_news)

            except Exception as error:
                return f"Erro ao buscar noticias confiaveis sobre IA: {error}\n"

        if not any([precisa_email, precisa_reuniao, precisa_slack, precisa_tarefa]):
            return ""

        try:
            eventos = await db.recent_events(limit=200)
        except Exception as error:
            return f"Erro ao buscar eventos no banco local do Heimdall: {error}\n"

        if not eventos:
            return "Nenhum evento foi encontrado no banco local do Heimdall.\n"

        fontes_desejadas: set[str] = set()

        if precisa_email:
            fontes_desejadas.add("email")

        if precisa_reuniao:
            fontes_desejadas.add("calendar")

        if precisa_slack:
            fontes_desejadas.add("slack")

        if precisa_tarefa:
            fontes_desejadas.add("task")

        eventos_filtrados = []

        for evento in eventos or []:
            source = valor_evento(evento, "source", "module", "type").lower()

            if not fontes_desejadas:
                eventos_filtrados.append(evento)
                continue

            if any(fonte in source for fonte in fontes_desejadas):
                eventos_filtrados.append(evento)

        if not eventos_filtrados:
            return "Não encontrei eventos compatíveis com a pergunta nos dados recentes do Heimdall.\n"

        linhas = [
            "Contexto operacional do Heimdall, do item mais recente para o mais antigo:",
            "Observação: o item 1 deve ser tratado como o mais recente entre os dados disponíveis.",
        ]

        for idx, evento in enumerate(eventos_filtrados[:12], start=1):
            meta = metadata_evento(evento)

            title = valor_evento(evento, "title", "subject", "name", default="Sem título")
            source = valor_evento(evento, "source", "module", "type", default="")
            category = valor_evento(evento, "category", "categoria", default="sem categoria")
            detail = valor_evento(evento, "detail", "details", "description", "body", "snippet", default="")
            score = valor_evento(evento, "score", "priority_score", default="")
            created_at = valor_evento(evento, "created_at", "date", "datetime", "timestamp", default="")

            if not detail:
                detail = str(
                    meta.get("detail")
                    or meta.get("description")
                    or meta.get("snippet")
                    or meta.get("body")
                    or meta.get("location")
                    or ""
                )

            linhas.append(
                f"{idx}. Título/assunto: {title}\n"
                f"   Fonte: {source}\n"
                f"   Categoria: {category}\n"
                f"   Score: {score}\n"
                f"   Data/hora registrada: {created_at}\n"
                f"   Detalhe disponível: {detail if detail else 'Sem detalhe adicional disponível.'}"
            )

        return "\n".join(linhas)

    def task_file_path() -> Path:
        return Path(deep_get(config, "agents.task_tracker.local_json_path", "data/tasks.json"))

    def load_local_tasks() -> list[dict[str, Any]]:
        path = task_file_path()
        path.parent.mkdir(parents=True, exist_ok=True)

        if not path.exists():
            path.write_text("[]", encoding="utf-8")

        try:
            data = json.loads(path.read_text(encoding="utf-8") or "[]")
            return data if isinstance(data, list) else []
        except Exception:
            return []

    def save_local_tasks(tasks: list[dict[str, Any]]) -> None:
        path = task_file_path()
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(tasks, ensure_ascii=False, indent=2), encoding="utf-8")

    BR_TZ = ZoneInfo("America/Sao_Paulo")
    MAX_FILE_TEXT_CHARS = 12000
    MAX_TOTAL_ATTACHMENT_CONTEXT_CHARS = 24000

    def task_upload_dir() -> Path:
        return Path(deep_get(config, "agents.task_tracker.upload_dir", "data/task_uploads"))

    def limpar_texto_extraido(text: str) -> str:
        text = str(text or "")
        text = text.replace("\x00", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def limitar_texto(text: str, limit: int = MAX_FILE_TEXT_CHARS) -> str:
        text = limpar_texto_extraido(text)

        if len(text) <= limit:
            return text

        return text[:limit].rstrip() + "\n\n[Texto cortado pelo Heimdall para não estourar o contexto da LLM.]"

    def ler_texto_simples(path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                return path.read_text(encoding=encoding, errors="ignore")
            except Exception:
                continue

        return ""

    def ler_csv(path: Path) -> str:
        raw = ler_texto_simples(path)

        if not raw:
            return ""

        try:
            sample = raw[:4096]
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except Exception:
            dialect = csv.excel

        output = io.StringIO()

        try:
            reader = csv.reader(io.StringIO(raw), dialect)

            for row_index, row in enumerate(reader):
                if row_index >= 120:
                    output.write("\n[CSV cortado após 120 linhas.]")
                    break

                output.write(" | ".join(str(cell).strip() for cell in row[:30]))
                output.write("\n")

            return output.getvalue()

        except Exception:
            return raw

    def ler_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
        except Exception as exc:
            return f"[Erro: biblioteca pypdf não instalada. Detalhe: {exc}]"

        textos: list[str] = []

        try:
            reader = PdfReader(str(path))

            for page_index, page in enumerate(reader.pages[:20], start=1):
                try:
                    page_text = page.extract_text() or ""
                except Exception:
                    page_text = ""

                if page_text.strip():
                    textos.append(f"--- Página {page_index} ---\n{page_text.strip()}")

            if not textos:
                return "[PDF sem texto extraível. Provavelmente é escaneado/imagem e exigiria OCR.]"

            return "\n\n".join(textos)

        except Exception as exc:
            return f"[Erro ao ler PDF: {exc}]"

    def ler_docx(path: Path) -> str:
        try:
            from docx import Document
        except Exception as exc:
            return f"[Erro: biblioteca python-docx não instalada. Detalhe: {exc}]"

        try:
            doc = Document(str(path))
            partes: list[str] = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    partes.append(text)

            for table_index, table in enumerate(doc.tables, start=1):
                partes.append(f"\n--- Tabela {table_index} ---")

                for row in table.rows[:80]:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells[:20]]
                    if any(cells):
                        partes.append(" | ".join(cells))

            return "\n".join(partes)

        except Exception as exc:
            return f"[Erro ao ler DOCX: {exc}]"

    def ler_xlsx(path: Path) -> str:
        try:
            from openpyxl import load_workbook
        except Exception as exc:
            return f"[Erro: biblioteca openpyxl não instalada. Detalhe: {exc}]"

        try:
            wb = load_workbook(str(path), read_only=True, data_only=True)
            partes: list[str] = []

            for ws in wb.worksheets[:8]:
                partes.append(f"\n--- Aba: {ws.title} ---")

                for row in ws.iter_rows(min_row=1, max_row=120, max_col=30, values_only=True):
                    values = ["" if value is None else str(value).strip() for value in row]

                    if any(values):
                        partes.append(" | ".join(values))

                if ws.max_row and ws.max_row > 120:
                    partes.append("[Aba cortada após 120 linhas.]")

            return "\n".join(partes)

        except Exception as exc:
            return f"[Erro ao ler XLSX/XLSM: {exc}]"

    def ler_xls(path: Path) -> str:
        try:
            import xlrd
        except Exception as exc:
            return f"[Erro: biblioteca xlrd não instalada. Detalhe: {exc}]"

        try:
            book = xlrd.open_workbook(str(path))
            partes: list[str] = []

            for sheet in book.sheets()[:8]:
                partes.append(f"\n--- Aba: {sheet.name} ---")

                max_rows = min(sheet.nrows, 120)
                max_cols = min(sheet.ncols, 30)

                for row_index in range(max_rows):
                    values = [
                        "" if sheet.cell_value(row_index, col_index) is None else str(sheet.cell_value(row_index, col_index)).strip()
                        for col_index in range(max_cols)
                    ]

                    if any(values):
                        partes.append(" | ".join(values))

                if sheet.nrows > 120:
                    partes.append("[Aba cortada após 120 linhas.]")

            return "\n".join(partes)

        except Exception as exc:
            return f"[Erro ao ler XLS: {exc}]"

    def extrair_texto_anexo(path: Path, original_name: str, content_type: str = "") -> dict[str, Any]:
        suffix = path.suffix.lower()
        content_type = str(content_type or "").lower()

        try:
            if suffix == ".pdf" or "pdf" in content_type:
                text = ler_pdf(path)

            elif suffix == ".docx":
                text = ler_docx(path)

            elif suffix == ".doc":
                text = (
                    "[Arquivo .doc antigo recebido. O Heimdall salva o arquivo, mas a leitura automática "
                    "confiável exige converter para .docx ou usar LibreOffice/antiword.]"
                )

            elif suffix in {".xlsx", ".xlsm"}:
                text = ler_xlsx(path)

            elif suffix == ".xls":
                text = ler_xls(path)

            elif suffix == ".csv":
                text = ler_csv(path)

            elif suffix in {".txt", ".md", ".json", ".log"}:
                text = ler_texto_simples(path)

            else:
                text = f"[Tipo de arquivo não suportado para leitura automática: {suffix or content_type or 'desconhecido'}]"

            text = limitar_texto(text)

            return {
                "ok": bool(text and not text.startswith("[Erro")),
                "text": text,
                "chars": len(text),
                "message": "Texto extraído." if text else "Nenhum texto extraído.",
            }

        except Exception as exc:
            return {
                "ok": False,
                "text": "",
                "chars": 0,
                "message": f"Erro ao extrair texto de {original_name}: {exc}",
            }

    def montar_contexto_anexos(attachments: list[dict[str, Any]]) -> str:
        if not attachments:
            return ""

        partes: list[str] = []

        for idx, attachment in enumerate(attachments, start=1):
            name = attachment.get("name", f"arquivo_{idx}")
            content_type = attachment.get("content_type", "")
            size_bytes = attachment.get("size_bytes", 0)
            extract_status = attachment.get("extract_status", "")
            text_preview = attachment.get("text_preview", "")

            partes.append(
                f"Arquivo {idx}: {name}\n"
                f"Tipo: {content_type or 'não informado'}\n"
                f"Tamanho: {size_bytes} bytes\n"
                f"Status da leitura: {extract_status or 'não informado'}"
            )

            if text_preview:
                partes.append(f"Conteúdo extraído:\n{text_preview}")
            else:
                partes.append("Conteúdo extraído: vazio ou não disponível.")

        context = "\n\n".join(partes)
        return limitar_texto(context, MAX_TOTAL_ATTACHMENT_CONTEXT_CHARS)

    def parse_datetime_safe(value: str | None) -> datetime | None:
        if not value:
            return None

        try:
            parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=BR_TZ)
            return parsed.astimezone(timezone.utc)
        except Exception:
            return None

    def fallback_relative_deadline(text: str) -> str:
        text_original = str(text or "").strip()
        text_norm = _normalize_text(text_original)

        now_local = datetime.now(BR_TZ)
        now_utc = datetime.now(timezone.utc)

        match = re.search(r"\bem\s+(\d+)\s*(minuto|minutos|min|mins|hora|horas|h|dia|dias)\b", text_norm)

        if match:
            amount = int(match.group(1))
            unit = match.group(2)

            if unit.startswith("min"):
                return (now_utc + timedelta(minutes=amount)).isoformat()

            if unit in {"h", "hora", "horas"}:
                return (now_utc + timedelta(hours=amount)).isoformat()

            if unit.startswith("dia"):
                return (now_utc + timedelta(days=amount)).isoformat()

        time_match = re.search(
            r"(?:as|às|a\s+partir\s+das)?\s*\b(\d{1,2})(?::|h)?(\d{2})?\b",
            text_norm,
        )

        hour: int | None = None
        minute: int = 0

        if time_match:
            possible_hour = int(time_match.group(1))
            possible_minute = int(time_match.group(2) or 0)

            if 0 <= possible_hour <= 23 and 0 <= possible_minute <= 59:
                hour = possible_hour
                minute = possible_minute

        if "hoje" in text_norm and hour is not None:
            due_local = now_local.replace(
                hour=hour,
                minute=minute,
                second=0,
                microsecond=0,
            )

            if due_local <= now_local:
                due_local = due_local + timedelta(days=1)

            return due_local.astimezone(timezone.utc).isoformat()

        if "amanha" in text_norm:
            if hour is None:
                hour = 9
                minute = 0

            due_local = (now_local + timedelta(days=1)).replace(
                hour=hour,
                minute=minute,
                second=0,
                microsecond=0,
            )

            return due_local.astimezone(timezone.utc).isoformat()

        date_match = re.search(r"\b(\d{1,2})/(\d{1,2})(?:/(\d{2,4}))?\b", text_norm)

        if date_match:
            day = int(date_match.group(1))
            month = int(date_match.group(2))
            year_raw = date_match.group(3)

            if year_raw:
                year = int(year_raw)
                if year < 100:
                    year += 2000
            else:
                year = now_local.year

            if hour is None:
                hour = 9
                minute = 0

            try:
                due_local = datetime(
                    year,
                    month,
                    day,
                    hour,
                    minute,
                    tzinfo=BR_TZ,
                )

                if due_local <= now_local and not year_raw:
                    due_local = due_local.replace(year=year + 1)

                return due_local.astimezone(timezone.utc).isoformat()

            except Exception:
                pass

        return ""

    async def inferir_tarefa_com_gemma(
        prompt: str,
        explicit_deadline: str = "",
        attachments_context: str = "",
    ) -> dict[str, Any]:
        now_local = datetime.now(BR_TZ).isoformat()

        system_prompt = (
            "Voce e o Agente de Tarefas do Heimdall. "
            "Sua funcao e transformar pedidos naturais e anexos em uma tarefa estruturada. "
            "Use o conteudo dos anexos para melhorar o titulo, lembrete, contexto e palavras-chave. "
            "Nao invente informacoes que nao estejam no pedido ou nos anexos. "
            "Responda apenas JSON valido, sem markdown. "
            "Campos obrigatorios: title, reminder_text, due_at, is_reminder, keywords, attachment_summary. "
            "due_at deve ser ISO 8601 com timezone. Se nao houver prazo, retorne string vazia. "
            "is_reminder deve ser true quando o usuario pedir lembrete, aviso, notificacao ou 'me lembre'. "
            "keywords deve ser uma lista curta de palavras relevantes. "
            "attachment_summary deve resumir em portugues o que foi lido nos anexos em ate 500 caracteres."
        )

        user_prompt = (
            f"Agora em Sao Paulo: {now_local}\n"
            f"Prazo informado manualmente, se houver: {explicit_deadline or ''}\n"
            f"Pedido do usuario: {prompt}\n\n"
            f"Conteudo lido dos anexos, se houver:\n{attachments_context or 'Nenhum anexo com texto extraido.'}"
        )

        payload = {
            "model": LMSTUDIO_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.1,
            "max_tokens": 700,
            "stream": False,
        }

        try:
            async with httpx.AsyncClient(timeout=min(LMSTUDIO_TIMEOUT, 45)) as client:
                response = await client.post(f"{LMSTUDIO_BASE_URL}/chat/completions", json=payload)

            response.raise_for_status()
            data = response.json()
            content = str(data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
            parsed = _extract_json_from_llm(content)

            return parsed if isinstance(parsed, dict) else {}

        except Exception:
            return {}

    async def salvar_anexos_tarefa(files: list[UploadFile] | None) -> list[dict[str, Any]]:
        saved: list[dict[str, Any]] = []

        if not files:
            return saved

        upload_dir = task_upload_dir()
        upload_dir.mkdir(parents=True, exist_ok=True)

        for file in files:
            original_name = Path(file.filename or "arquivo").name
            safe_name = re.sub(r"[^a-zA-Z0-9._-]+", "_", original_name).strip("_") or "arquivo"
            stored_name = f"{uuid.uuid4().hex}_{safe_name}"
            target = upload_dir / stored_name

            content = await file.read()

            if len(content) > 25 * 1024 * 1024:
                saved.append(
                    {
                        "name": original_name,
                        "stored_name": "",
                        "path": "",
                        "content_type": file.content_type or "",
                        "size_bytes": len(content),
                        "extract_status": "Arquivo ignorado: maior que 25 MB.",
                        "text_preview": "",
                        "extracted_chars": 0,
                        "read_ok": False,
                    }
                )
                continue

            target.write_bytes(content)

            extraction = extrair_texto_anexo(
                path=target,
                original_name=original_name,
                content_type=file.content_type or "",
            )

            saved.append(
                {
                    "name": original_name,
                    "stored_name": stored_name,
                    "path": str(target),
                    "content_type": file.content_type or "",
                    "size_bytes": len(content),
                    "extract_status": extraction.get("message", ""),
                    "text_preview": extraction.get("text", ""),
                    "extracted_chars": extraction.get("chars", 0),
                    "read_ok": extraction.get("ok", False),
                }
            )

        return saved

    async def publicar_evento_tarefa(task: dict[str, Any], title: str, detail: str, score: int = 7) -> int:
        event = AgentEvent(
            source="task",
            title=title,
            detail=detail,
            score=score,
            urgency="media",
            category="tarefa",
            external_id=f"local:{task.get('id')}",
            metadata=task,
        )

        event_id = await db.save_event(event)
        await bus.publish({"type": "event", "event": {**event.model_dump(), "id": event_id}})
        return event_id

    async def criar_tarefa_agente(
        prompt: str,
        explicit_deadline: str = "",
        files: list[UploadFile] | None = None,
        source: str = "task_agent",
    ) -> dict[str, Any]:
        prompt = str(prompt or "").strip()

        if not prompt:
            return {"ok": False, "message": "Texto da tarefa vazio."}

        attachments = await salvar_anexos_tarefa(files)
        attachments_context = montar_contexto_anexos(attachments)

        inferred = await inferir_tarefa_com_gemma(
            prompt=prompt,
            explicit_deadline=explicit_deadline,
            attachments_context=attachments_context,
        )

        manual_deadline_dt = parse_datetime_safe(explicit_deadline)
        inferred_deadline_dt = parse_datetime_safe(str(inferred.get("due_at") or ""))
        fallback_deadline = fallback_relative_deadline(prompt)
        fallback_deadline_dt = parse_datetime_safe(fallback_deadline)

        now_utc = datetime.now(timezone.utc)

        # Se a IA/Gemma devolver uma data no passado, ignora.
        if inferred_deadline_dt and inferred_deadline_dt < now_utc - timedelta(seconds=30):
            inferred_deadline_dt = None

        # Se o fallback também ficar no passado, ignora.
        if fallback_deadline_dt and fallback_deadline_dt < now_utc - timedelta(seconds=30):
            fallback_deadline_dt = None

        # Campo manual tem prioridade, mas evita gravar data vencida sem querer.
        if manual_deadline_dt:
            if manual_deadline_dt < now_utc - timedelta(seconds=30):
                return {
                    "ok": False,
                    "message": "A data/hora informada já passou. Informe um horário futuro.",
                }

            deadline = manual_deadline_dt.isoformat()

        elif fallback_deadline_dt:
            deadline = fallback_deadline_dt.isoformat()

        elif inferred_deadline_dt:
            deadline = inferred_deadline_dt.isoformat()

        else:
            deadline = ""

        title = str(inferred.get("title") or prompt).strip()[:180]
        reminder_text = str(inferred.get("reminder_text") or title).strip()[:280]
        attachment_summary = str(inferred.get("attachment_summary") or "").strip()[:700]

        prompt_norm = _normalize_text(prompt)
        is_reminder = bool(inferred.get("is_reminder")) or any(
            term in prompt_norm
            for term in ["lembre", "lembrar", "lembrete", "me avise", "avisar", "notifique", "notificacao"]
        )

        task_id = f"heimdall-{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:6]}"

        keywords = inferred.get("keywords")
        if not isinstance(keywords, list):
            keywords = [word for word in title.lower().split()[:6] if len(word) > 3]

        task = {
            "id": task_id,
            "title": title,
            "reminder_text": reminder_text,
            "status": "open",
            "task_type": "reminder" if is_reminder else "task",
            "deadline": deadline,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "source": source,
            "keywords": keywords[:8],
            "attachments": attachments,
            "attachment_summary": attachment_summary,
            "notified_at": "",
            "completed_at": "",
        }

        tasks = load_local_tasks()
        tasks.append(task)
        save_local_tasks(tasks)

        detail_parts = ["Criada pelo Agente de Tarefas com Gemma local."]

        if attachment_summary:
            detail_parts.append(f"Resumo dos anexos: {attachment_summary}")

        if attachments:
            ok_count = sum(1 for item in attachments if item.get("read_ok"))
            detail_parts.append(f"Anexos recebidos: {len(attachments)}. Lidos com sucesso: {ok_count}.")

        event_id = await publicar_evento_tarefa(
            task,
            title=f"Nova tarefa: {task['title']}",
            detail=" ".join(detail_parts),
            score=7 if is_reminder else 6,
        )

        await db.log("INFO", "task", f"Tarefa criada: {task['title']}")

        return {"ok": True, "task": task, "event_id": event_id}

    async def task_agent_loop() -> None:
        while True:
            try:
                now = datetime.now(timezone.utc)
                tasks = load_local_tasks()
                changed = False

                for task in tasks:
                    if str(task.get("status") or "open").lower() not in {"open", "aberta"}:
                        continue

                    if str(task.get("task_type") or "") != "reminder":
                        continue

                    if task.get("notified_at"):
                        continue

                    deadline_dt = parse_datetime_safe(task.get("deadline"))
                    if not deadline_dt:
                        continue

                    if deadline_dt <= now:
                        task["notified_at"] = now.isoformat()
                        task["completed_at"] = now.isoformat()
                        task["status"] = "done"
                        changed = True

                        reminder_text = str(
                            task.get("reminder_text")
                            or task.get("title")
                            or "Lembrete concluido."
                        )

                        await publicar_evento_tarefa(
                            task,
                            title=f"Lembrete: {task.get('title')}",
                            detail=reminder_text,
                            score=8,
                        )

                        await bus.publish(
                            {
                                "type": "task_interaction",
                                "kind": "task_due",
                                "text": f"Rodrigo, lembrete: {reminder_text}",
                                "task": task,
                            }
                        )

                        await db.log("INFO", "task", f"Lembrete disparado e concluido: {task.get('title')}")

                if changed:
                    save_local_tasks(tasks)

            except asyncio.CancelledError:
                raise

            except Exception as exc:
                await db.log("ERROR", "task", f"Erro no agente de tarefas: {exc}")

            await asyncio.sleep(15)

    async def criar_tarefa_local(pergunta: str) -> dict[str, Any]:
        result = await criar_tarefa_agente(
            prompt=pergunta,
            explicit_deadline="",
            files=None,
            source="local_llm",
        )

        return {
            "action": "create_task",
            "task": result.get("task"),
            "event_id": result.get("event_id"),
            "ok": result.get("ok", False),
            "message": result.get("message", ""),
        }

    async def executar_acao_operacional(pergunta: str) -> dict[str, Any] | None:
        pergunta_norm = _normalize_text(pergunta)

        if any(term in pergunta_norm for term in ["crie tarefa", "criar tarefa", "crie uma tarefa", "nova tarefa", "adicionar tarefa", "adicione uma tarefa"]):
            return await criar_tarefa_local(pergunta)

        if any(term in pergunta_norm for term in ["noticia", "noticias", "inteligencia artificial", "ia do mercado", "ia hoje"]):
            return {"action": "fetch_ai_news", "news": await buscar_noticias_ia(limit=12)}

        if any(term in pergunta_norm for term in ["kpi", "indicador", "status do sistema", "resumo do sistema"]):
            return {"action": "read_kpis", "kpis": await db.kpis()}

        source_map = {
            "email": "email",
            "e-mail": "email",
            "reuniao": "calendar",
            "agenda": "calendar",
            "calendario": "calendar",
            "slack": "slack",
            "tarefa": "task",
            "tarefas": "task",
        }

        for term, source in source_map.items():
            if term in pergunta_norm and any(cmd in pergunta_norm for cmd in ["liste", "listar", "mostre", "mostrar", "ver", "consultar", "ultimos", "ultimas"]):
                return {
                    "action": "read_events",
                    "source": source,
                    "events": await db.recent_events(limit=15, source=source),
                }

        return None

    def parse_event_datetime(value: Any) -> datetime:
        raw = str(value or "").strip()

        if not raw:
            return datetime.now(timezone.utc)

        try:
            dt = datetime.fromisoformat(raw.replace("Z", "+00:00"))
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt.astimezone(timezone.utc)
        except Exception:
            return datetime.now(timezone.utc)

    def is_today_local(value: Any) -> bool:
        dt = parse_event_datetime(value).astimezone(BR_TZ)
        now = datetime.now(BR_TZ)
        return dt.date() == now.date()

    def build_gmail_url(event: dict[str, Any]) -> str:
        metadata = event.get("metadata") or {}
        thread_id = str(metadata.get("thread") or metadata.get("thread_id") or "").strip()
        message_id = str(metadata.get("message_id") or event.get("external_id") or "").strip()

        if thread_id:
            return f"https://mail.google.com/mail/u/0/#inbox/{thread_id}"

        if message_id:
            cleaned = message_id.strip("<>")
            return f"https://mail.google.com/mail/u/0/#search/rfc822msgid%3A{cleaned}"

        return "https://mail.google.com/mail/u/0/#inbox"

    async def chamar_gemma_json(
        *,
        system_prompt: str,
        user_prompt: str,
        max_tokens: int = 900,
        temperature: float = 0.15,
    ) -> dict[str, Any]:
        payload = {
            "model": LMSTUDIO_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": False,
        }

        async with httpx.AsyncClient(timeout=min(LMSTUDIO_TIMEOUT, 60)) as client:
            response = await client.post(f"{LMSTUDIO_BASE_URL}/chat/completions", json=payload)

        response.raise_for_status()
        data = response.json()
        content = str(data.get("choices", [{}])[0].get("message", {}).get("content") or "{}")
        parsed = _extract_json_from_llm(content)
        return parsed if isinstance(parsed, dict) else {}

    def fallback_email_executive(events: list[dict[str, Any]]) -> dict[str, Any]:
        if not events:
            return {
                "ok": True,
                "has_email": False,
                "message": "Nenhum email de hoje encontrado.",
                "checked_at": datetime.now(timezone.utc).isoformat(),
            }

        best = sorted(
            events,
            key=lambda event: (
                float(event.get("score") or 0),
                parse_event_datetime(event.get("created_at")).timestamp(),
            ),
            reverse=True,
        )[0]
        metadata = best.get("metadata") or {}
        sender = metadata.get("sender") or metadata.get("from") or ""
        subject = best.get("title") or "(sem assunto)"

        return {
            "ok": True,
            "has_email": True,
            "event_id": best.get("id"),
            "subject": subject,
            "sender": sender,
            "importance_score": float(best.get("score") or 0),
            "category": best.get("category") or "geral",
            "urgency": best.get("urgency") or "media",
            "summary": best.get("detail") or "",
            "reason": "Fallback: maior score entre os emails de hoje.",
            "recommended_action": "Abrir e avaliar resposta conforme prioridade.",
            "voice_alert": f"Rodrigo, o email mais importante agora e de {sender or 'remetente nao identificado'}. Assunto: {subject}.",
            "gmail_url": build_gmail_url(best),
            "checked_at": datetime.now(timezone.utc).isoformat(),
        }

    async def analyze_email_executive(force_notify: bool = False) -> dict[str, Any]:
        events = await db.recent_events(limit=300, source="email")
        today_emails = [event for event in events if is_today_local(event.get("created_at"))]

        if not today_emails:
            state = fallback_email_executive([])
            await db.set_state("email_executive_state", state)
            return state

        payload_items = []
        for event in today_emails[:40]:
            metadata = event.get("metadata") or {}
            payload_items.append(
                {
                    "id": event.get("id"),
                    "sender": metadata.get("sender") or metadata.get("from") or "",
                    "subject": event.get("title") or "",
                    "summary": event.get("detail") or "",
                    "score": event.get("score"),
                    "urgency": event.get("urgency"),
                    "category": event.get("category"),
                    "created_at": event.get("created_at"),
                    "reason": metadata.get("reason") or "",
                }
            )

        system_prompt = (
            "Voce e o Agente Executivo de Emails do Heimdall. "
            "Atue como secretario executivo senior, especialista em assuntos financeiros, administrativos, contratos, prazos, riscos e operacao. "
            "Escolha o email mais importante do dia para o usuario. "
            "Responda apenas JSON valido, sem markdown."
        )
        user_prompt = (
            "Analise os emails abaixo e retorne JSON com os campos: "
            "event_id, subject, sender, importance_score de 0 a 10, category, urgency, summary, reason, recommended_action, voice_alert.\n"
            "Priorize: impacto financeiro, aprovacao, prazo hoje, risco juridico, pagamento, erro operacional, diretoria, cliente importante e bloqueios.\n\n"
            f"Emails de hoje:\n{json.dumps(payload_items, ensure_ascii=False)}"
        )

        try:
            analyzed = await chamar_gemma_json(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                max_tokens=900,
                temperature=0.1,
            )
        except Exception as exc:
            await db.log("WARNING", "email", f"Agente executivo usou fallback: {exc}")
            analyzed = {}

        if not analyzed:
            state = fallback_email_executive(today_emails)
        else:
            selected_id = str(analyzed.get("event_id") or "")
            selected = next((event for event in today_emails if str(event.get("id")) == selected_id), None)

            if not selected:
                state = fallback_email_executive(today_emails)
            else:
                metadata = selected.get("metadata") or {}
                state = {
                    "ok": True,
                    "has_email": True,
                    "event_id": selected.get("id"),
                    "subject": str(analyzed.get("subject") or selected.get("title") or "(sem assunto)"),
                    "sender": str(analyzed.get("sender") or metadata.get("sender") or metadata.get("from") or ""),
                    "importance_score": max(0, min(10, float(analyzed.get("importance_score") or selected.get("score") or 0))),
                    "category": str(analyzed.get("category") or selected.get("category") or "geral"),
                    "urgency": str(analyzed.get("urgency") or selected.get("urgency") or "media"),
                    "summary": str(analyzed.get("summary") or selected.get("detail") or "")[:1200],
                    "reason": str(analyzed.get("reason") or "")[:700],
                    "recommended_action": str(analyzed.get("recommended_action") or "Abrir e responder se necessario.")[:500],
                    "voice_alert": str(analyzed.get("voice_alert") or "")[:500],
                    "gmail_url": build_gmail_url(selected),
                    "checked_at": datetime.now(timezone.utc).isoformat(),
                }

                if not state["voice_alert"]:
                    state["voice_alert"] = f"Rodrigo, email importante de {state['sender']}. Assunto: {state['subject']}."

        await db.set_state("email_executive_state", state)

        signature = f"{state.get('event_id')}|{state.get('importance_score')}|{state.get('subject')}"
        last_signature = await db.get_state("email_executive_last_signature", "")

        should_notify = bool(state.get("has_email")) and (force_notify or signature != last_signature) and float(state.get("importance_score") or 0) >= 7

        if should_notify:
            await db.set_state("email_executive_last_signature", signature)
            await bus.publish({"type": "email_executive", "state": state})
            await db.log("INFO", "email", f"Email executivo destacado: {state.get('subject')}")

        return state

    async def email_executive_agent_loop() -> None:
        await asyncio.sleep(8)

        while True:
            try:
                await analyze_email_executive()
            except asyncio.CancelledError:
                raise
            except Exception as exc:
                await db.log("ERROR", "email", f"Erro no Agente Executivo de Emails: {exc}")

            await asyncio.sleep(3600)

    def collect_project_code_inventory() -> dict[str, Any]:
        root_path = Path.cwd()
        extensions = {
            "python": [".py"],
            "javascript": [".js", ".jsx", ".ts", ".tsx"],
            "sql": [".sql"],
            "html_css": [".html", ".css"],
            "r": [".r", ".R"],
        }
        ignored_parts = {".venv", "__pycache__", ".git", "node_modules", "data", "logs"}
        inventory: dict[str, Any] = {key: {"files": [], "total": 0} for key in extensions}

        for path in root_path.rglob("*"):
            if not path.is_file():
                continue
            if any(part in ignored_parts for part in path.parts):
                continue
            suffix = path.suffix
            for key, suffixes in extensions.items():
                if suffix in suffixes:
                    rel = str(path.relative_to(root_path))
                    inventory[key]["files"].append(rel)
                    inventory[key]["total"] += 1

        for value in inventory.values():
            value["files"] = value["files"][:80]

        return inventory

    def read_code_sample(file_path: str, limit: int = 16000) -> str:
        root_path = Path.cwd().resolve()
        if not file_path:
            return ""
        target = (root_path / file_path).resolve()
        if not target.exists() or not target.is_file() or root_path not in target.parents:
            return ""
        if target.stat().st_size > 512 * 1024:
            return "[Arquivo grande demais para leitura direta no painel.]"
        return limitar_texto(ler_texto_simples(target), limit)

    async def analyze_coders(payload: CoderAnalyzeRequest) -> dict[str, Any]:
        inventory = collect_project_code_inventory()
        code_sample = str(payload.code or "").strip()[:16000] or read_code_sample(payload.file_path)
        prompt = str(payload.prompt or "").strip() or "Analise riscos, bugs provaveis e pontos de melhoria."

        def fallback_coder_agents(reason: str = "") -> dict[str, Any]:
            language_totals = {
                key: int(value.get("total") or 0)
                for key, value in inventory.items()
                if isinstance(value, dict)
            }
            return {
                "summary": "Analise estrutural em fallback: a LLM local nao retornou JSON valido, mas o inventario foi carregado.",
                "risks": [
                    "A resposta da LLM local precisa estar em JSON para alimentar os cinco cards automaticamente.",
                    "O painel esta em modo auditoria: ele recomenda achados, mas nao altera arquivos sozinho.",
                ],
                "recommendations": [
                    "Confirme no LM Studio se o modelo Gemma esta carregado e se aceita instrucoes de JSON estrito.",
                    "Ao analisar um arquivo especifico, informe o caminho relativo no campo do Coders.",
                ],
                "agents": [
                    {
                        "name": "Python Sentinel",
                        "language": "python",
                        "status": "fallback",
                        "finding": f"{language_totals.get('python', 0)} arquivo(s) Python mapeado(s).",
                        "risk": reason[:220],
                        "recommendation": "Priorizar app.py, agentes e servicos locais nas proximas analises.",
                    },
                    {
                        "name": "JavaScript Watcher",
                        "language": "javascript",
                        "status": "fallback",
                        "finding": f"{language_totals.get('javascript', 0)} arquivo(s) JavaScript/TypeScript mapeado(s).",
                        "risk": "",
                        "recommendation": "Separar scripts grandes em modulos quando a base crescer.",
                    },
                    {
                        "name": "SQL Auditor",
                        "language": "sql",
                        "status": "fallback",
                        "finding": f"{language_totals.get('sql', 0)} arquivo(s) SQL mapeado(s).",
                        "risk": "",
                        "recommendation": "Validar queries e migracoes quando forem adicionadas ao projeto.",
                    },
                    {
                        "name": "Interface Guardian",
                        "language": "html_css",
                        "status": "fallback",
                        "finding": f"{language_totals.get('html_css', 0)} arquivo(s) HTML/CSS mapeado(s).",
                        "risk": "Template concentrado pode dificultar manutencao de UI.",
                        "recommendation": "Extrair CSS/JS do dashboard para arquivos separados em uma etapa futura.",
                    },
                    {
                        "name": "R Analyst",
                        "language": "r",
                        "status": "fallback",
                        "finding": f"{language_totals.get('r', 0)} arquivo(s) R mapeado(s).",
                        "risk": "",
                        "recommendation": "Ativar este agente quando rotinas estatisticas em R forem anexadas.",
                    },
                ],
            }

        system_prompt = (
            "Voce e o Conclave Coders do Heimdall: cinco agentes especialistas em Python, JavaScript, SQL, HTML/CSS e R. "
            "Analise rigorosamente a estrutura. "
            "OBRIGATORIO: retorne EXCLUSIVAMENTE um objeto JSON valido, sem markdown ou texto extra, no seguinte formato exato:\n"
            '{"agents": [{"name": "Agent", "language": "python", "status": "ok", "finding": "...", "risk": "...", "recommendation": "..."}]}'
        )
        user_prompt = (
            f"Inventario do projeto Heimdall:\n{json.dumps(inventory, ensure_ascii=False)}\n\n"
            f"Pedido do usuario: {prompt}\n"
            f"Linguagem/escopo informado: {payload.language}\n"
            f"Arquivo informado: {payload.file_path}\n\n"
            f"Codigo anexado/lido, se houver:\n{code_sample or 'Sem codigo especifico; analise a estrutura pelo inventario.'}"
        )

        try:
            result = await chamar_gemma_json(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                max_tokens=1600,
                temperature=0.18,
            )
        except Exception as exc:
            result = fallback_coder_agents(str(exc))

        agents = result.get("agents") if isinstance(result, dict) else []
        if not isinstance(agents, list):
            agents = []

        if not agents:
            result = fallback_coder_agents("Sem agentes validos na resposta da LLM.")
            agents = result["agents"]

        return {
            "ok": True,
            "checked_at": datetime.now(timezone.utc).isoformat(),
            "inventory": inventory,
            "summary": str(result.get("summary") or "Analise concluida.")[:1200],
            "risks": result.get("risks") if isinstance(result.get("risks"), list) else [],
            "recommendations": result.get("recommendations") if isinstance(result.get("recommendations"), list) else [],
            "agents": agents[:5],
        }
    
    async def coders_conclave_loop():
        await asyncio.sleep(30)
        while True:
            try:
                payload = CoderAnalyzeRequest(
                    prompt="Auditoria automatica de rotina pelo Coders Conclave.", 
                    language="auto", 
                    file_path="src"
                )
                result = await analyze_coders(payload)
                await db.log("INFO", "coders", "Auditoria automatica do Coders Conclave concluida.")
                await bus.publish({"type": "coders_update", "data": result})
            except Exception as e:
                await db.log("ERROR", "coders", f"Erro no loop do Coders Conclave: {e}")
            await asyncio.sleep(3600)

    @app.on_event("startup")
    async def iniciar_agente_tarefas():
        app.state.task_agent_task = asyncio.create_task(task_agent_loop())
        app.state.email_executive_task = asyncio.create_task(email_executive_agent_loop())
        app.state.coders_conclave_task = asyncio.create_task(coders_conclave_loop())
        await db.log("INFO", "task", "Agente de tarefas iniciado.")
        await db.log("INFO", "email", "Agente Executivo de Emails iniciado.")
        await db.log("INFO", "coders", "Agente Coders Conclave iniciado.")

    @app.on_event("shutdown")
    async def parar_agente_tarefas():
        task = getattr(app.state, "task_agent_task", None)
        if task:
            task.cancel()

        email_task = getattr(app.state, "email_executive_task", None)
        if email_task:
            email_task.cancel()

        coders_task = getattr(app.state, "coders_conclave_task", None)
        if coders_task:
            coders_task.cancel()

    @app.get("/", response_class=HTMLResponse)
    async def root(request: Request):
        return templates.TemplateResponse(
            request=request,
            name="dashboard.html",
            context={
                "title": config.get("dashboard", {}).get("title", "Heimdall"),
                "voice_language": settings.voice_language,
            },
        )

    @app.get("/dashboard", response_class=HTMLResponse)
    async def dashboard(request: Request):
        return templates.TemplateResponse(
            request=request,
            name="dashboard.html",
            context={
                "title": config.get("dashboard", {}).get("title", "Heimdall"),
                "voice_language": settings.voice_language,
            },
        )

    @app.get("/api/kpis")
    async def api_kpis():
        return await db.kpis()

    @app.get("/api/events")
    async def api_events(limit: int = 30):
        return await db.recent_events(limit=limit)

    @app.get("/api/logs")
    async def api_logs(limit: int = 100):
        return await db.recent_logs(limit=limit)

    @app.get("/api/activity")
    async def api_activity():
        return await db.weekly_activity()

    @app.get("/api/tasks")
    async def api_tasks():
        tasks = load_local_tasks()
        tasks.sort(key=lambda task: str(task.get("created_at") or ""), reverse=True)
        return {"ok": True, "items": tasks}

    @app.post("/api/tasks")
    async def api_create_task(payload: TaskCreateRequest):
        return await criar_tarefa_agente(
            prompt=payload.title,
            explicit_deadline=payload.deadline or "",
            files=None,
            source="dashboard",
        )
    
    @app.post("/api/tasks/agent")
    async def api_create_task_agent(
        prompt: str = Form(...),
        deadline: str = Form(""),
        files: list[UploadFile] | None = File(default=None),
    ):
        return await criar_tarefa_agente(
            prompt=prompt,
            explicit_deadline=deadline,
            files=files,
            source="dashboard_upload",
        )

    @app.patch("/api/tasks/{task_id}")
    async def api_update_task(task_id: str, payload: TaskUpdateRequest):
        tasks = load_local_tasks()
        updated: dict[str, Any] | None = None

        for task in tasks:
            if str(task.get("id")) != task_id:
                continue

            if payload.status is not None:
                task["status"] = payload.status
            if payload.title is not None and payload.title.strip():
                task["title"] = payload.title.strip()[:180]
            if payload.deadline is not None:
                task["deadline"] = payload.deadline

            task["updated_at"] = datetime.now(timezone.utc).isoformat()
            updated = task
            break

        if not updated:
            return {"ok": False, "message": "Tarefa nao encontrada."}

        save_local_tasks(tasks)
        await db.log("INFO", "task", f"Tarefa atualizada: {updated.get('title')}")
        return {"ok": True, "task": updated}
    
    @app.delete("/api/tasks/{task_id}")
    async def api_delete_task(task_id: str):
        tasks = load_local_tasks()

        deleted_task: dict[str, Any] | None = None
        remaining_tasks: list[dict[str, Any]] = []

        for task in tasks:
            if str(task.get("id")) == str(task_id):
                deleted_task = task
                continue

            remaining_tasks.append(task)

        if not deleted_task:
            return {
                "ok": False,
                "message": "Lembrete nao encontrado.",
            }

        removed_files = 0

        try:
            upload_root = task_upload_dir().resolve()

            for attachment in deleted_task.get("attachments", []) or []:
                raw_path = str(attachment.get("path") or "").strip()

                if not raw_path:
                    continue

                attachment_path = Path(raw_path)

                try:
                    resolved_path = attachment_path.resolve()
                except Exception:
                    continue

                if (
                    resolved_path.exists()
                    and resolved_path.is_file()
                    and upload_root in resolved_path.parents
                ):
                    resolved_path.unlink()
                    removed_files += 1

        except Exception as exc:
            await db.log("WARNING", "task", f"Lembrete excluido, mas houve falha ao remover anexos: {exc}")

        save_local_tasks(remaining_tasks)

        title = str(deleted_task.get("title") or "Lembrete sem titulo")

        await db.log("INFO", "task", f"Lembrete excluido: {title}")

        await bus.publish(
            {
                "type": "task_interaction",
                "kind": "task_deleted",
                "text": f"Lembrete excluido: {title}",
                "task": deleted_task,
            }
        )

        return {
            "ok": True,
            "message": "Lembrete excluido com sucesso.",
            "deleted_id": task_id,
            "removed_files": removed_files,
        }

    @app.get("/api/news/ai")
    async def api_ai_news(limit: int = 12):
        return await buscar_noticias_ia(limit=limit)

    @app.get("/api/email-executive/current")
    async def api_email_executive_current():
        state = await db.get_state("email_executive_state", None)

        if state:
            return state

        return await analyze_email_executive(force_notify=False)

    @app.post("/api/email-executive/run")
    async def api_email_executive_run():
        return await analyze_email_executive(force_notify=True)

    @app.get("/api/coders/inventory")
    async def api_coders_inventory():
        return {
            "ok": True,
            "checked_at": datetime.now(timezone.utc).isoformat(),
            "inventory": collect_project_code_inventory(),
        }

    @app.post("/api/coders/analyze")
    async def api_coders_analyze(payload: CoderAnalyzeRequest):
        return await analyze_coders(payload)

    # ========================================================
    # ROTAS DA LLM LOCAL
    # ========================================================

    @app.get("/api/local-llm/status")
    async def local_llm_status():
        try:
            models = await testar_conexao_lmstudio()

            return {
                "ok": True,
                "message": "LM Studio conectado.",
                "model_configurado": LMSTUDIO_MODEL,
                "models": models,
            }

        except Exception as exc:
            return {
                "ok": False,
                "message": (
                    "Não foi possível conectar ao LM Studio. "
                    "Verifique se o LM Studio está aberto, se o modelo está carregado "
                    "e se o Local Server está ativo em http://localhost:1234."
                ),
                "erro": str(exc),
            }

    @app.post("/api/local-llm/chat")
    async def local_llm_chat(payload: LocalLLMRequest):
        try:
            contexto_operacional = await montar_contexto_operacional(payload.pergunta)
            acao = await executar_acao_operacional(payload.pergunta)

            if acao:
                contexto_operacional = (
                    f"{contexto_operacional}\n\n"
                    "Resultado de acao operacional executada pelo Heimdall:\n"
                    f"{json.dumps(acao, ensure_ascii=False, default=str)}"
                ).strip()

            resposta = await perguntar_llm_local(
                payload.pergunta,
                contexto_operacional=contexto_operacional,
                historico=payload.historico,
                imagens=payload.imagens,
            )

            return {
                "ok": True,
                "resposta": resposta,
                "contexto_usado": bool(contexto_operacional),
                "imagens_recebidas": len(payload.imagens),
                "acao": acao,
            }

        except Exception as exc:
            return {
                "ok": False,
                "resposta": "",
                "message": "Erro ao consultar a LLM local.",
                "erro": str(exc),
            }

    @app.post("/api/local-llm/action")
    async def local_llm_action(payload: LocalActionRequest):
        try:
            acao = await executar_acao_operacional(payload.pergunta)
            contexto_operacional = await montar_contexto_operacional(payload.pergunta)

            if acao:
                contexto_operacional = (
                    f"{contexto_operacional}\n\n"
                    "Resultado de acao operacional executada pelo Heimdall:\n"
                    f"{json.dumps(acao, ensure_ascii=False, default=str)}"
                ).strip()

            resposta = await perguntar_llm_local(
                payload.pergunta,
                contexto_operacional=contexto_operacional,
                historico=payload.historico,
                imagens=payload.imagens,
            )

            return {
                "ok": True,
                "resposta": resposta,
                "acao": acao,
                "contexto_usado": bool(contexto_operacional),
                "imagens_recebidas": len(payload.imagens),
            }

        except Exception as exc:
            return {
                "ok": False,
                "resposta": "",
                "message": "Erro ao executar acao com a LLM local.",
                "erro": str(exc),
            }

    @app.post("/api/voice-command")
    async def api_voice_command(payload: VoicePayload):
        answer = await handle_voice_command(payload.text, db)

        await db.log(
            "INFO",
            "voice",
            f"Comando de voz: {payload.text} -> {answer['intent']}",
        )

        await bus.publish(
            {
                "type": "voice",
                "text": payload.text,
                "answer": answer,
            }
        )

        return answer

    @app.post("/api/language-tutor/chat")
    async def language_tutor_chat(payload: LanguageTutorRequest):
        language_code = payload.language if payload.language in LANGUAGE_TUTOR_OPTIONS else "en-US"
        language = LANGUAGE_TUTOR_OPTIONS[language_code]
        user_text = payload.message.strip()

        if not user_text:
            return {"ok": False, "reply": "", "message": "Mensagem vazia."}

        messages: list[dict[str, str]] = [
            {
                "role": "system",
                "content": (
                    "Voce e o agente Professor de Idiomas do Heimdall, rodando localmente com Gemma. "
                    f"O idioma principal da aula e {language['name']} ({language['native']}). "
                    "O usuario pode escrever em portugues, no idioma alvo ou misturar os dois. "
                    "Se ele escrever em portugues, traduza a ideia para o idioma alvo e continue a aula. "
                    "Se ele escrever no idioma alvo, responda naturalmente nesse idioma e corrija erros. "
                    "Mantenha a conversa fluida, com frases curtas, didaticas e praticas. "
                    "Nao transforme tudo em explicacao longa. A prioridade e conversacao. "
                    "Sempre siga este formato:\n"
                    "Resposta: <resposta natural no idioma alvo>\n"
                    "Traducao: <tradução curta em portugues brasileiro, quando ajudar>\n"
                    "Correcao: <correcao objetiva em portugues brasileiro, ou 'Sem correcao importante.'>\n"
                    "Proxima pergunta: <pergunta curta no idioma alvo para manter a conversa>"
                ),
            }
        ]

        for item in payload.history[-10:]:
            role = item.role if item.role in {"user", "assistant"} else "user"
            content = str(item.content or "").strip()
            if content:
                messages.append({"role": role, "content": content})

        messages.append({"role": "user", "content": user_text})

        try:
            async with httpx.AsyncClient(timeout=min(LMSTUDIO_TIMEOUT, 60)) as client:
                response = await client.post(
                    f"{LMSTUDIO_BASE_URL}/chat/completions",
                    json={
                        "model": LMSTUDIO_MODEL,
                        "messages": messages,
                        "temperature": 0.45,
                        "top_p": 0.9,
                        "max_tokens": 550,
                        "stream": False,
                    },
                )
            response.raise_for_status()
            data = response.json()
            reply = str(data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()

            if not reply:
                reply = (
                    f"Resposta: Let's continue practicing {language['native']}.\n"
                    "Correcao: A LLM local retornou vazio nesta rodada.\n"
                    "Proxima pergunta: What would you like to talk about?"
                )

            await db.log("INFO", "voice", f"Language tutor {language_code}: {user_text[:120]}")

            return {
                "ok": True,
                "language": language_code,
                "reply": reply,
            }

        except Exception as exc:
            return {
                "ok": False,
                "language": language_code,
                "reply": "",
                "message": "Erro ao consultar o professor de idiomas local.",
                "erro": str(exc),
            }

    @app.websocket("/ws")
    async def ws_endpoint(websocket: WebSocket):
        await websocket.accept()
        await bus.register(websocket)

        try:
            await websocket.send_json(
                {
                    "type": "hello",
                    "message": "Heimdall conectado",
                }
            )

            while True:
                # Mantém a conexão aberta.
                # O front-end pode mandar mensagens/ping por aqui.
                await websocket.receive_text()

        except WebSocketDisconnect:
            await bus.unregister(websocket)

        except Exception:
            await bus.unregister(websocket)

    return app
